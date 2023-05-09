from docx import Document

document = Document()

paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
document.add_heading('The REAL meaning of the universe')

table = document.add_table(rows=2, cols=2)
cell = table.cell(0, 1)
cell.text = 'parrot, possibly dead'

document.save('test.docx')